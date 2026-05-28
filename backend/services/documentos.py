# backend/services/documentos.py
# Word: Arial 11pt, espaçamento 1.15, margens 2cm
# PDF: idêntico ao Word com timbre FOA como fundo

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import io, re
from pathlib import Path

A4_W_EMU = 7560000
A4_H_EMU = 10692000
TIMBRE_PATH = str(Path(__file__).parent.parent.parent / "modelos" / "timbre.png")

BOLD_PATTERNS = [
    r'FUNDAÇÃO OSWALDO ARANHA\s*[–\-]\s*FOA',
    r'FUNDACAO OSWALDO ARANHA\s*[–\-]\s*FOA',
    r'FUNDAÇÃO OSWALDO ARANHA',
    r'FUNDACAO OSWALDO ARANHA',
    r'EXECUÇÃO DE TÍTULO EXTRAJUDICIAL',
    r'EXECUÇÃO FISCAL',
    r'AÇÃO DE COBRANÇA',
    r'AÇÃO MONITÓRIA',
    r'AÇÃO ORDINÁRIA',
    r'exclusivamente em nome do Assessor Jurídico da Fundação, DENYS RIBEIRO FURTUNATO - OAB/RJ 164\.024',
]


def _field(p, code):
    r = p.add_run()
    r.font.name = 'Arial'; r.font.size = Pt(10); r.bold = True
    e1 = OxmlElement('w:fldChar'); e1.set(qn('w:fldCharType'), 'begin'); r._r.append(e1)
    i  = OxmlElement('w:instrText'); i.text = f' {code} '; r._r.append(i)
    e2 = OxmlElement('w:fldChar'); e2.set(qn('w:fldCharType'), 'end'); r._r.append(e2)


def _add_page_number(paragraph):
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after  = Pt(0)
    r0 = paragraph.add_run("Página "); r0.font.name='Arial'; r0.font.size=Pt(10); r0.bold=True
    _field(paragraph, 'PAGE')
    r1 = paragraph.add_run(" de ");   r1.font.name='Arial'; r1.font.size=Pt(10); r1.bold=True
    _field(paragraph, 'NUMPAGES')


def _run(p, texto, bold=False):
    r = p.add_run(texto)
    r.font.name = 'Arial'; r.font.size = Pt(11); r.bold = bold
    return r


def _pf(p, indent=None):
    pf = p.paragraph_format
    pf.space_before      = Pt(0)
    pf.space_after       = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = Pt(11 * 1.15)
    pf.first_line_indent = indent
    return pf


def _vazio(doc):
    p = doc.add_paragraph()
    _pf(p)
    return p


def _add_watermark(doc, image_path):
    if not Path(image_path).exists():
        return
    section = doc.sections[0]
    header  = section.header
    for p in header.paragraphs:
        p.clear()
    hp = header.paragraphs[0]
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after  = Pt(0)
    run = hp.add_run()
    run.add_picture(image_path, width=Cm(21), height=Cm(29.7))
    drawing_elem = run._r.find(qn('w:drawing'))
    inline_elem  = drawing_elem.find(
        '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
    blip = inline_elem.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
    rId  = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
    run._r.remove(drawing_elem)
    anchor_xml = f'''<w:drawing
        xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
        xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"
        xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
      <wp:anchor distT="0" distB="0" distL="0" distR="0"
                 simplePos="0" relativeHeight="251658240" behindDoc="1"
                 locked="1" layoutInCell="1" allowOverlap="0">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionH>
        <wp:positionV relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionV>
        <wp:extent cx="{A4_W_EMU}" cy="{A4_H_EMU}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:wrapNone/>
        <wp:docPr id="1001" name="Timbre"/>
        <wp:cNvGraphicFramePr><a:graphicFrameLocks noChangeAspect="1"/></wp:cNvGraphicFramePr>
        <a:graphic>
          <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
            <pic:pic>
              <pic:nvPicPr>
                <pic:cNvPr id="1002" name="Timbre"/>
                <pic:cNvPicPr><a:picLocks noChangeAspect="1"/></pic:cNvPicPr>
              </pic:nvPicPr>
              <pic:blipFill>
                <a:blip r:embed="{rId}"/>
                <a:stretch><a:fillRect/></a:stretch>
              </pic:blipFill>
              <pic:spPr>
                <a:xfrm><a:off x="0" y="0"/><a:ext cx="{A4_W_EMU}" cy="{A4_H_EMU}"/></a:xfrm>
                <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
              </pic:spPr>
            </pic:pic>
          </a:graphicData>
        </a:graphic>
      </wp:anchor>
    </w:drawing>'''
    run._r.append(etree.fromstring(anchor_xml))


def _escrever_com_negritos(p, texto, parte_contraria=None):
    padroes = list(BOLD_PATTERNS)
    if parte_contraria and parte_contraria.strip():
        padroes.append(re.escape(parte_contraria.strip()))
    regex  = '(' + '|'.join(padroes) + ')'
    partes = re.split(regex, texto, flags=re.IGNORECASE)
    for parte in partes:
        if not parte:
            continue
        negrito = any(re.fullmatch(pat, parte, flags=re.IGNORECASE) for pat in padroes)
        _run(p, parte, bold=negrito)


def gerar_docx(numero_processo: str, conteudo: str,
               parte_contraria: str = None) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.page_width    = Cm(21)
        section.page_height   = Cm(29.7)
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    style.paragraph_format.space_before      = Pt(0)
    style.paragraph_format.space_after       = Pt(0)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing      = Pt(11 * 1.15)

    _add_watermark(doc, TIMBRE_PATH)

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    _add_page_number(fp)

    linhas = [l.strip() for l in conteudo.split('\n')]
    idx_endereco = next((i for i,l in enumerate(linhas) if l.startswith('Ao Douto') or l.startswith('Douto Juízo')), None)
    idx_processo = next((i for i,l in enumerate(linhas) if l.startswith('Processo')), None)
    idx_data     = next((i for i,l in enumerate(linhas) if l.startswith('Volta Redonda,')), None)

    i = 0
    while i < len(linhas):
        ls = linhas[i]

        if i == idx_endereco:
            p = doc.add_paragraph(); _pf(p); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _run(p, ls, bold=True)
            for _ in range(5): _vazio(doc)
            i += 1; continue

        if i == idx_processo:
            p = doc.add_paragraph(); _pf(p); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _run(p, ls, bold=True)
            _vazio(doc)
            i += 1; continue

        if 'exclusivamente em nome do Assessor' in ls or ls.startswith('Requer por fim') or ls.startswith('Por fim, requer'):
            _vazio(doc)
            p = doc.add_paragraph(); _pf(p, indent=Cm(2.5)); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _run(p, ls, bold=True)
            _vazio(doc)
            i += 1; continue

        if ls.startswith('P. Defer') or ls.startswith('Pede Defer') or ls.startswith('Nestes Termos'):
            _vazio(doc)
            p = doc.add_paragraph(); _pf(p, indent=Cm(2.5)); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _run(p, ls)
            i += 1; continue

        if ls.startswith('Volta Redonda,'):
            p = doc.add_paragraph(); _pf(p, indent=Cm(2.5)); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _run(p, ls)
            for _ in range(4): _vazio(doc)
            i += 1; continue

        if any(x in ls for x in ['Advogado(a)', 'OAB/RJ', 'OAB/']):
            p = doc.add_paragraph(); _pf(p); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(p, ls, bold=True)
            i += 1; continue

        if idx_data is not None and i > idx_data and ls:
            palavras = ls.split()
            if (len(palavras) <= 5 and ls[0].isupper() and
                not any(x in ls for x in ['OAB','Advogado','Fundação','FOA','Requer','Diante','FUNDAÇÃO','FUNDACAO'])):
                p = doc.add_paragraph(); _pf(p); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _run(p, ls, bold=True)
                i += 1; continue

        if not ls:
            _vazio(doc)
            i += 1; continue

        p = doc.add_paragraph(); _pf(p, indent=Cm(2.5)); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _escrever_com_negritos(p, ls, parte_contraria)
        _vazio(doc)
        i += 1

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def gerar_pdf(numero_processo: str, conteudo: str,
              parte_contraria: str = None) -> bytes:
    """
    Converte o .docx para PDF usando python-docx2pdf ou
    gera via Word COM no Windows. Fallback: ReportLab simples.
    """
    # Tenta converter via LibreOffice ou docx2pdf
    try:
        import subprocess, tempfile, os
        docx_bytes = gerar_docx(numero_processo, conteudo, parte_contraria)

        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = os.path.join(tmpdir, 'peticao.docx')
            pdf_path  = os.path.join(tmpdir, 'peticao.pdf')

            with open(docx_path, 'wb') as f:
                f.write(docx_bytes)

            # Tenta LibreOffice (Linux/Mac)
            for soffice in ['soffice', 'libreoffice']:
                try:
                    result = subprocess.run(
                        [soffice, '--headless', '--convert-to', 'pdf',
                         '--outdir', tmpdir, docx_path],
                        capture_output=True, timeout=30
                    )
                    if result.returncode == 0 and os.path.exists(pdf_path):
                        with open(pdf_path, 'rb') as f:
                            return f.read()
                except Exception:
                    continue

            # Tenta docx2pdf (Windows com Word instalado)
            try:
                from docx2pdf import convert
                convert(docx_path, pdf_path)
                if os.path.exists(pdf_path):
                    with open(pdf_path, 'rb') as f:
                        return f.read()
            except Exception:
                pass

    except Exception:
        pass

    # Fallback ReportLab com Arial via fonte registrada
    return _gerar_pdf_reportlab(numero_processo, conteudo, parte_contraria)


def _gerar_pdf_reportlab(numero_processo: str, conteudo: str,
                          parte_contraria: str = None) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm, mm
    from reportlab.platypus import BaseDocTemplate, PageTemplate, Frame, Paragraph, Spacer
    from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from pathlib import Path
    import re

    buffer = io.BytesIO()
    W, H   = A4

    # Registra Arial do Windows
    _F, _FB = 'Helvetica', 'Helvetica-Bold'
    for reg, bold in [
        (r'C:\Windows\Fonts\arial.ttf',   r'C:\Windows\Fonts\arialbd.ttf'),
        (r'C:\Windows\Fonts\Arial.ttf',   r'C:\Windows\Fonts\Arial Bold.ttf'),
        ('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
         '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf'),
    ]:
        if Path(reg).exists():
            try:
                pdfmetrics.registerFont(TTFont('Arial',      reg))
                if Path(bold).exists():
                    pdfmetrics.registerFont(TTFont('Arial-Bold', bold))
                else:
                    pdfmetrics.registerFont(TTFont('Arial-Bold', reg))
                _F, _FB = 'Arial', 'Arial-Bold'
                break
            except Exception:
                continue

    TIMBRE = TIMBRE_PATH
    LH = 11 * 1.15  # leading em pts

    def _fundo(canvas, doc):
        if Path(TIMBRE).exists():
            canvas.saveState()
            canvas.drawImage(TIMBRE, 0, 0, width=W, height=H,
                             preserveAspectRatio=False, mask='auto')
            canvas.restoreState()

    def _esc(t):
        return t.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')

    def _bold_html(texto):
        padroes = list(BOLD_PATTERNS)
        if parte_contraria:
            padroes.append(re.escape(parte_contraria.strip()))
        resultado = _esc(texto)
        for pat in padroes:
            resultado = re.sub(f'({pat})', r'<b>\1</b>', resultado, flags=re.IGNORECASE)
        return resultado

    def E(align=TA_LEFT, indent=0, bold=False):
        return ParagraphStyle('x',
            fontName=_FB if bold else _F,
            fontSize=11, alignment=align,
            firstLineIndent=indent,
            spaceAfter=0, spaceBefore=0, leading=LH)

    M = 2*cm
    frame = Frame(M, M, W-2*M, H-2*M, leftPadding=0, rightPadding=0,
                  topPadding=0, bottomPadding=0)
    tmpl  = PageTemplate(id='foa', frames=[frame], onPage=_fundo)
    doc   = BaseDocTemplate(buffer, pagesize=A4,
                            leftMargin=M, rightMargin=M,
                            topMargin=M, bottomMargin=M)
    doc.addPageTemplates([tmpl])

    VAZIO = Spacer(1, LH)
    linhas = [l.strip() for l in conteudo.split('\n')]
    idx_end  = next((i for i,l in enumerate(linhas) if l.startswith('Ao Douto') or l.startswith('Douto Juízo')), None)
    idx_proc = next((i for i,l in enumerate(linhas) if l.startswith('Processo')), None)
    idx_data = next((i for i,l in enumerate(linhas) if l.startswith('Volta Redonda,')), None)

    els = []
    i = 0
    while i < len(linhas):
        ls = linhas[i]

        if i == idx_end:
            els.append(Paragraph(f'<b>{_esc(ls)}</b>', E(TA_LEFT, 0, True)))
            for _ in range(5): els.append(VAZIO)
            i += 1; continue

        if i == idx_proc:
            els.append(Paragraph(f'<b>{_esc(ls)}</b>', E(TA_LEFT, 0, True)))
            els.append(VAZIO)
            i += 1; continue

        if 'exclusivamente em nome do Assessor' in ls or ls.startswith('Requer por fim') or ls.startswith('Por fim, requer'):
            els.append(VAZIO)
            els.append(Paragraph(f'<b>{_esc(ls)}</b>', E(TA_JUSTIFY, 2.5*cm, True)))
            els.append(VAZIO)
            i += 1; continue

        if ls.startswith('P. Defer') or ls.startswith('Pede Defer'):
            els.append(VAZIO)
            els.append(Paragraph(_esc(ls), E(TA_LEFT, 2.5*cm)))
            i += 1; continue

        if ls.startswith('Volta Redonda,'):
            els.append(Paragraph(_esc(ls), E(TA_LEFT, 2.5*cm)))
            for _ in range(4): els.append(VAZIO)
            i += 1; continue

        if any(x in ls for x in ['Advogado(a)', 'OAB/RJ', 'OAB/']):
            els.append(Paragraph(f'<b>{_esc(ls)}</b>', E(TA_CENTER, 0, True)))
            i += 1; continue

        if idx_data is not None and i > idx_data and ls:
            palavras = ls.split()
            if (len(palavras) <= 5 and ls[0].isupper() and
                not any(x in ls for x in ['OAB','Advogado','Requer','FUNDACAO','FUNDAÇÃO'])):
                els.append(Paragraph(f'<b>{_esc(ls)}</b>', E(TA_CENTER, 0, True)))
                i += 1; continue

        if not ls:
            els.append(VAZIO)
            i += 1; continue

        els.append(Paragraph(_bold_html(ls), E(TA_JUSTIFY, 2.5*cm)))
        els.append(VAZIO)
        i += 1

    doc.build(els)
    return buffer.getvalue()